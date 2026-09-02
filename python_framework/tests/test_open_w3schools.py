import os
from playwright.sync_api import sync_playwright
from pathlib import Path
import time


def test_open_w3schools():
    headless_env = os.getenv("HEADLESS", "1")
    headless = headless_env != "0"

    base_dir = Path(__file__).resolve().parents[1]
    screenshots_dir = base_dir / "screenshots"
    screenshots_dir.mkdir(exist_ok=True)

    steps = []

    with sync_playwright() as p:
        browser = p.chromium.launch(headless=headless)
        context = browser.new_context(viewport={"width": 1920, "height": 1080})
        page = context.new_page()

        # Step 1: Open target URL
        page.goto("https://www.w3schools.com/")
        page.wait_for_load_state("networkidle")
        s1 = screenshots_dir / "step1_home.png"
        page.screenshot(path=str(s1), full_page=True)
        steps.append(("Open W3Schools homepage", s1.name))

        # Step 2: Click Python tab
        clicked = False
        try:
            page.click("a[href^='/python']", timeout=5000)
            clicked = True
        except Exception:
            # Fallback: click link by text
            try:
                page.click("text=Python", timeout=5000)
                clicked = True
            except Exception:
                clicked = False

        if clicked:
            try:
                page.wait_for_url("**/python*", timeout=15000)
            except Exception:
                try:
                    page.wait_for_load_state("domcontentloaded", timeout=15000)
                except Exception:
                    pass
            time.sleep(1)
            s2 = screenshots_dir / "step2_python_page.png"
            page.screenshot(path=str(s2), full_page=True)
            steps.append(("Python tutorial page", s2.name))
        else:
            s2 = screenshots_dir / "step2_python_page_missing.png"
            page.screenshot(path=str(s2), full_page=True)
            steps.append(("Python tab not found - captured homepage", s2.name))

        # Step 3: Click 'Try it Yourself' and switch to the new page or same tab
        tryit = None
        try:
            # Try to open in a new page
            with context.expect_page(timeout=15000) as new_page_info:
                page.locator("a:has-text('Try it Yourself')").first.click()
            tryit = new_page_info.value
        except Exception:
            try:
                # Maybe it opens in same tab
                page.locator("a:has-text('Try it Yourself')").first.click(timeout=7000)
                tryit = page
            except Exception:
                # Fallback: search for link and navigate
                href = page.get_attribute("a:has-text('Try it Yourself')", "href")
                if href:
                    tryit = context.new_page()
                    tryit.goto(href)

        if not tryit:
            raise RuntimeError("Could not open Try it Yourself page")

        try:
            tryit.wait_for_load_state("domcontentloaded", timeout=15000)
        except Exception:
            time.sleep(1)
        s3 = screenshots_dir / "step3_tryit_open.png"
        tryit.screenshot(path=str(s3), full_page=False)
        steps.append(("Try it Yourself editor opened", s3.name))

        # Step 4: Replace code in editor and run
        code = 'print("Hello, Playwright Automation")'
        # Try common textarea selector first
        editor_filled = False
        try:
            ta = tryit.wait_for_selector("textarea#textareaCode, textarea", timeout=7000)
            ta.fill(code)
            editor_filled = True
        except Exception:
            # Fallback: try CodeMirror or set textarea via JS
            try:
                filled = tryit.evaluate(
                    '''() => {
                    const ta = document.querySelector('textarea');
                    if (ta) { ta.value = 'print("Hello, Playwright Automation")'; return true; }
                    if (window.editor && typeof window.editor.setValue === 'function') { window.editor.setValue('print("Hello, Playwright Automation")'); return true; }
                    const cm = document.querySelector('.CodeMirror');
                    if (cm && cm.CodeMirror) { cm.CodeMirror.setValue('print("Hello, Playwright Automation")'); return true; }
                    return false;
                    }''')
                editor_filled = bool(filled)
            except Exception:
                editor_filled = False

        s4_before = screenshots_dir / "step4_editor_before.png"
        tryit.screenshot(path=str(s4_before), full_page=False)
        steps.append(("Editor before run", s4_before.name))

        # Click Run
        clicked_run = False
        try:
            tryit.click("button#runbtn", timeout=7000)
            clicked_run = True
        except Exception:
            try:
                tryit.click("text=Run", timeout=7000)
                clicked_run = True
            except Exception:
                clicked_run = False

        time.sleep(2)
        # Step 4b: capture result
        s5 = screenshots_dir / "step5_result.png"
        try:
            # try to screenshot iframe result
            iframe_el = tryit.query_selector("iframe#iframeResult, iframe")
            if iframe_el:
                # capture viewport-size screenshot of iframe
                box = iframe_el.bounding_box()
                if box:
                    img = tryit.screenshot(path=str(s5), full_page=False)
                else:
                    iframe_el.screenshot(path=str(s5))
            else:
                tryit.screenshot(path=str(s5), full_page=False)
            steps.append(("Result after Run", s5.name))
        except Exception:
            tryit.screenshot(path=str(s5), full_page=False)
            steps.append(("Result (fallback)", s5.name))

        # Optionally verify output text inside result iframe
        output_ok = False
        try:
            frame = None
            for f in tryit.frames:
                if 'result' in (f.name or '').lower() or 'iframeResult' in (f.name or ''):
                    frame = f
                    break
            if not frame and len(tryit.frames) > 1:
                frame = tryit.frames[-1]
            if frame:
                body = frame.evaluate('() => document.body.innerText')
                if 'Hello, Playwright Automation' in (body or ''):
                    output_ok = True
        except Exception:
            output_ok = False

        # Step 5: Generate a Word document report with screenshots
        try:
            from docx import Document
            from docx.shared import Inches

            doc = Document()
            doc.add_heading('W3Schools TryIt Automation Report', level=1)
            for desc, img in steps:
                doc.add_heading(desc, level=2)
                img_path = screenshots_dir / img
                if img_path.exists():
                    doc.add_picture(str(img_path), width=Inches(6))
            doc.add_paragraph(f"Output verification: {'PASS' if output_ok else 'FAIL'}")
            doc_path = base_dir / 'w3schools_tryit_report.docx'
            doc.save(doc_path)
        except Exception:
            # Fallback to HTML report if python-docx not available
            report_path = base_dir / "w3schools_tryit_report.html"
            with open(report_path, "w", encoding="utf-8") as fh:
                fh.write("<html><head><meta charset='utf-8'><title>W3Schools TryIt Report</title></head><body>")
                fh.write(f"<h1>W3Schools TryIt Automation Report</h1>")
                for desc, img in steps:
                    fh.write(f"<h2>{desc}</h2>")
                    fh.write(f"<img src=\"screenshots/{img}\" style=\"max-width:100%;border:1px solid #ccc\"><br/>")
                fh.write(f"<h3>Output verification: {'PASS' if output_ok else 'FAIL'}</h3>")
                fh.write("</body></html>")

        browser.close()

