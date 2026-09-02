import os
from pathlib import Path
from playwright.sync_api import sync_playwright
import time


def test_open_swayam():
    headless_env = os.getenv("HEADLESS", "1")
    headless = headless_env != "0"

    base_dir = Path(__file__).resolve().parents[1]
    screenshots_dir = base_dir / "screenshots"
    screenshots_dir.mkdir(exist_ok=True)

    with sync_playwright() as p:
        browser = p.chromium.launch(headless=headless)
        context = browser.new_context(viewport={"width": 1920, "height": 1080})
        page = context.new_page()

        # Open home
        page.goto("https://swayam.gov.in/")
        page.wait_for_load_state("domcontentloaded")
        shot = screenshots_dir / "swayam_home.png"
        page.screenshot(path=str(shot), full_page=False)

        # Step: navigate to Courses / About Courses
        clicked = False
        course_selectors = [
            "text=All Courses",
            "text=Courses",
            "a[href*='course']",
            "text=Course",
        ]
        for sel in course_selectors:
            try:
                page.click(sel, timeout=5000)
                clicked = True
                break
            except Exception:
                continue

        # If not clicked, try opening the menu then clicking
        if not clicked:
            try:
                # try common menu selector
                page.click("button[aria-label='Toggle navigation']", timeout=3000)
                for sel in course_selectors:
                    try:
                        page.click(sel, timeout=5000)
                        clicked = True
                        break
                    except Exception:
                        continue
            except Exception:
                pass

        time.sleep(1)
        s_courses = screenshots_dir / "swayam_courses_nav.png"
        page.screenshot(path=str(s_courses), full_page=False)

        # Step: ensure we're on All Courses page — if not, try finding an 'All Courses' link and clicking
        try:
            page.wait_for_url("**/course/**", timeout=7000)
        except Exception:
            try:
                # find link by text
                page.click("text=All Courses", timeout=5000)
            except Exception:
                pass

        time.sleep(1)
        s_allcourses = screenshots_dir / "swayam_all_courses.png"
        page.screenshot(path=str(s_allcourses), full_page=False)

        # Step: locate search input and type 'java'
        search_filled = False
        search_selectors = [
            "input[type='search']",
            "input[placeholder*='Search']",
            "input[aria-label*='Search']",
            "input[name*='search']",
        ]
        for sel in search_selectors:
            try:
                inp = page.wait_for_selector(sel, timeout=4000)
                inp.fill("java")
                inp.press("Enter")
                search_filled = True
                break
            except Exception:
                continue

        # fallback: try a generic text input
        if not search_filled:
            try:
                inp = page.locator("input").first
                inp.fill("java")
                inp.press("Enter")
                search_filled = True
            except Exception:
                search_filled = False

        time.sleep(2)
        s_search = screenshots_dir / "swayam_search_java.png"
        page.screenshot(path=str(s_search), full_page=False)

        # capture results area if possible
        try:
            results = page.locator("text=java", timeout=3000)
            if results.count() > 0:
                s_result = screenshots_dir / "swayam_search_result.png"
                page.screenshot(path=str(s_result), full_page=False)
        except Exception:
            s_result = screenshots_dir / "swayam_search_result_fallback.png"
            page.screenshot(path=str(s_result), full_page=False)

        # Build Word document report
        try:
            from docx import Document
            from docx.shared import Inches

            doc = Document()
            doc.add_heading('Swayam Course Search Report', level=1)
            doc.add_heading('Home', level=2)
            doc.add_picture(str(shot), width=Inches(6))
            doc.add_heading('Courses navigation', level=2)
            doc.add_picture(str(s_courses), width=Inches(6))
            doc.add_heading('All Courses', level=2)
            doc.add_picture(str(s_allcourses), width=Inches(6))
            doc.add_heading('Search "java" results', level=2)
            doc.add_picture(str(s_search), width=Inches(6))
            if 's_result' in locals() and (screenshots_dir / s_result.name).exists():
                doc.add_paragraph('Result snapshot:')
                doc.add_picture(str(screenshots_dir / s_result.name), width=Inches(6))
            doc_path = base_dir / 'swayam_course_search_report.docx'
            doc.save(doc_path)
        except Exception:
            # Fallback HTML
            report_path = base_dir / "swayam_course_search_report.html"
            with open(report_path, "w", encoding="utf-8") as fh:
                fh.write("<html><body>")
                fh.write("<h1>Swayam Course Search Report</h1>")
                fh.write(f"<h2>Home</h2><img src=\"screenshots/{shot.name}\" style=\"max-width:100%\"><br/>")
                fh.write(f"<h2>Courses nav</h2><img src=\"screenshots/{s_courses.name}\" style=\"max-width:100%\"><br/>")
                fh.write(f"<h2>All courses</h2><img src=\"screenshots/{s_allcourses.name}\" style=\"max-width:100%\"><br/>")
                fh.write(f"<h2>Search</h2><img src=\"screenshots/{s_search.name}\" style=\"max-width:100%\"><br/>")
                if 's_result' in locals():
                    fh.write(f"<h2>Result</h2><img src=\"screenshots/{s_result.name}\" style=\"max-width:100%\"><br/>")
                fh.write("</body></html>")

        browser.close()
